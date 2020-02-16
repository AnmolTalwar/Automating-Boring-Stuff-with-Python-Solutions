import docx,os,re


os.chdir("C:\\Python\\Udemy\\Assignment\\Invitation")

doc = docx.Document('Invite.docx')

names = open('List.txt')
list = names.readlines()

for name in list:
    doc.add_paragraph('It would be a pleasure to have the company of').style = 'Normal'
    name = re.sub("\n","",name)
    doc.add_paragraph(name).style = 'Normal'
    doc.add_paragraph('at 11010 Memory Lane on the Evening of').style = 'Normal'
    doc.add_paragraph('April 1st').style = 'Normal'
    Para = doc.add_paragraph('at 7 o\'clock')
    Para.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    Para.style='Normal'

doc.save('Invite.docx')
print("Invitations Printed!!")
